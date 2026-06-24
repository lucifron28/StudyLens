import os
import django
from django.core.files.base import ContentFile
import fitz  # PyMuPDF
from docx import Document

# Set up Django environment
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings')
django.setup()

from django.contrib.auth.models import User
from learning.models import Subject, Module

def generate_pdf(filename, title, text):
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), title, fontsize=20)
    page.insert_textbox(fitz.Rect(50, 80, 550, 800), text, fontsize=12)
    doc.save(filename)
    doc.close()
    return filename

def generate_docx(filename, title, text):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(text)
    doc.save(filename)
    return filename

def populate():
    # Ensure demo_files directory exists
    os.makedirs('demo_files', exist_ok=True)

    # Get the user "ron" or the first available user
    owner = User.objects.filter(username="ron").first() or User.objects.first()
    if not owner:
        print("No users found in the database. Please create a user first.")
        return

    # 1. C Programming Fundamentals
    subj_c, _ = Subject.objects.get_or_create(
        owner=owner,
        title="C Programming Fundamentals",
        defaults={
            "description": "A comprehensive guide to understanding the C programming language, its structure, memory management, and pointers.",
            "color": "#3B82F6"
        }
    )
    
    Module.objects.get_or_create(
        owner=owner,
        subject=subj_c,
        title="1. Introduction and Data Types",
        defaults={
            "description": "Learn about the basic structure of a C program and various data types.",
            "content_type": "markdown",
            "markdown_content": "# Introduction to C\n\nC is a powerful, general-purpose programming language developed in the 1970s. It is known for its efficiency and low-level access to memory.\n\n## Basic Structure\nA C program typically begins with preprocessor directives and the `main()` function.\n```c\n#include <stdio.h>\nint main() {\n    printf(\"Hello World\");\n    return 0;\n}\n```\n\n## Data Types\n- `int`: Integers (e.g., 10, -5)\n- `float`: Floating point numbers (e.g., 3.14)\n- `char`: Single characters (e.g., 'A')\n- `double`: Double precision floating point numbers"
        }
    )

    Module.objects.get_or_create(
        owner=owner,
        subject=subj_c,
        title="2. Pointers & Memory Management",
        defaults={
            "description": "Master one of C's most powerful features: pointers and manual memory allocation.",
            "content_type": "markdown",
            "markdown_content": "# Pointers in C\n\nPointers store the memory address of another variable. They are essential for dynamic memory allocation and efficient array manipulation.\n\n```c\nint num = 10;\nint *ptr = &num; // ptr stores the address of num\nprintf(\"Value: %d\", *ptr); // Dereferencing\n```\n\n## Memory Management\nC allows manual allocation and deallocation of memory using functions like `malloc()`, `calloc()`, and `free()` from `<stdlib.h>`.\n```c\nint *arr = (int*)malloc(5 * sizeof(int));\nfree(arr);\n```"
        }
    )

    # 2. Object Oriented Programming
    subj_oop, _ = Subject.objects.get_or_create(
        owner=owner,
        title="Object Oriented Programming",
        defaults={
            "description": "Learn the core concepts of OOP including encapsulation, inheritance, and polymorphism.",
            "color": "#10B981"
        }
    )

    Module.objects.get_or_create(
        owner=owner,
        subject=subj_oop,
        title="Classes and Objects",
        defaults={
            "description": "Understanding the blueprints of OOP.",
            "content_type": "markdown",
            "markdown_content": "# Classes and Objects\n\n## Classes\nA class is a blueprint for creating objects. It encapsulates data for the object and methods to manipulate that data.\n\n## Objects\nAn object is an instance of a class. When a class is defined, no memory is allocated but when it is instantiated (i.e. an object is created) memory is allocated.\n\n### Example (Java)\n```java\nclass Car {\n    String color;\n    void drive() {\n        System.out.println(\"Driving...\");\n    }\n}\nCar myCar = new Car();\n```"
        }
    )

    # 3. HTTP vs HTTPS
    subj_http, _ = Subject.objects.get_or_create(
        owner=owner,
        title="HTTP vs HTTPS",
        defaults={
            "description": "Understand web protocols, data transmission, and how SSL/TLS encryption secures the web.",
            "color": "#F59E0B"
        }
    )

    Module.objects.get_or_create(
        owner=owner,
        subject=subj_http,
        title="Understanding the Protocols",
        defaults={
            "description": "The difference between HTTP and HTTPS.",
            "content_type": "markdown",
            "markdown_content": "# HTTP vs HTTPS\n\n## HTTP (Hypertext Transfer Protocol)\nHTTP is the foundation of data communication for the World Wide Web. Data sent over HTTP is in plain text, meaning anyone intercepting the traffic can read it.\n\n## HTTPS (HTTP Secure)\nHTTPS uses SSL/TLS to encrypt HTTP requests and responses. This ensures that attackers cannot eavesdrop or tamper with the data.\n\n### Key Differences\n1. **Security**: HTTPS is encrypted, HTTP is not.\n2. **Port**: HTTP uses port 80, HTTPS uses port 443.\n3. **Certificates**: HTTPS requires an SSL/TLS certificate."
        }
    )

    # 4. System Design
    subj_sd, _ = Subject.objects.get_or_create(
        owner=owner,
        title="System Design",
        defaults={
            "description": "Learn how to architect large scale, distributed, and highly available software systems.",
            "color": "#8B5CF6"
        }
    )

    Module.objects.get_or_create(
        owner=owner,
        subject=subj_sd,
        title="Load Balancing & Scalability",
        defaults={
            "description": "Techniques for scaling web applications.",
            "content_type": "markdown",
            "markdown_content": "# Scalability\n\n## Vertical vs Horizontal Scaling\n- **Vertical Scaling (Scale Up)**: Adding more power (CPU, RAM) to your existing machine.\n- **Horizontal Scaling (Scale Out)**: Adding more machines into your pool of resources.\n\n## Load Balancers\nA load balancer distributes incoming network traffic across multiple servers. This ensures no single server bears too much demand, improving responsiveness and availability.\n\nAlgorithms include:\n- Round Robin\n- Least Connections\n- IP Hash"
        }
    )

    # Generate sample documents
    print("Generating demo PDF and DOCX files...")
    
    pdf_text = "This is a sample document about Artificial Intelligence.\nAI is transforming the modern world, impacting everything from healthcare to autonomous vehicles.\nMachine learning models, particularly deep neural networks, have driven recent breakthroughs."
    generate_pdf("demo_files/ai_overview.pdf", "Artificial Intelligence Overview", pdf_text)
    
    docx_text = "This is a sample Word document regarding Project Management.\nEffective project management requires clear communication, realistic timelines, and proper resource allocation. Agile methodologies have become increasingly popular in software development."
    generate_docx("demo_files/project_management.docx", "Project Management Guide", docx_text)

    print("Database populated successfully and demo files created in backend/demo_files/")

if __name__ == '__main__':
    populate()
